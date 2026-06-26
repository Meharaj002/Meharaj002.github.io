"""Builder for the ATS-clean cv.docx (single column, standard fonts, real text,
no images/tables/columns). Edit the strings below and re-run with the project
venv python to regenerate cv.docx:  ../backend/.venv/Scripts/python.exe _gen_cv_docx.py
"""
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

MUTED = RGBColor(0x55, 0x55, 0x55)
RULE = "C9C9C9"

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.5)
    s.left_margin = s.right_margin = Inches(0.6)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.line_spacing = 1.12

TEXT_W = Inches(7.3)


def md_runs(p, text):
    bold = False
    for part in text.split("**"):
        if part:
            p.add_run(part).bold = bold
        bold = not bold


def heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), RULE)
    pbdr.append(bottom)
    pPr.append(pbdr)


def role(pos, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(TEXT_W, WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(pos).bold = True
    if date:
        r = p.add_run("\t" + date)
        r.font.size = Pt(9.5)
        r.font.color.rgb = MUTED


def para(text, justify=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    md_runs(p, text)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    md_runs(p, text)


h = doc.add_paragraph()
h.paragraph_format.space_after = Pt(0)
r = h.add_run("MEHARAJ SAMI SIAM")
r.bold = True
r.font.size = Pt(20)
para("Applied AI Engineer — Agentic AI, LLM, RAG & MCP").runs[0].bold = True
c = doc.add_paragraph()
c.paragraph_format.space_after = Pt(6)
cr = c.add_run(
    "Dhaka, Bangladesh  |  mssiam12@gmail.com  |  +880 1796-537490  |  "
    "linkedin.com/in/meharaj-sami-siam-a4b730213  |  github.com/Meharaj002  |  "
    "dorpon-eyewear.netlify.app"
)
cr.font.size = Pt(9.3)
cr.font.color.rgb = MUTED

heading("Professional Summary")
para(
    "Applied AI Engineer building production, agentic AI and Large Language Model (LLM) "
    "applications — three-tier LLM routing, Retrieval-Augmented Generation (RAG) pipelines, "
    "Model Context Protocol (MCP) integrations, and LLM fine-tuning on AWS Bedrock (Anthropic "
    "Claude). Strong full-stack (Django, Next.js, Vue) and REST API / SDK engineering "
    "foundation, with hands-on machine learning, prompt engineering, and scalable cloud "
    "deployment (AWS). Has delivered 13+ projects end-to-end, from architecture and design "
    "through production deployment, including a multi-tenant AI platform hardened to enterprise "
    "security standards and 6 production SDKs.",
    justify=True,
)

heading("Core Competencies")
para(
    "Agentic AI & MCP  |  Large Language Models (LLMs)  |  Retrieval-Augmented Generation (RAG)"
    "  |  LLM Fine-Tuning  |  Prompt Engineering  |  Machine Learning  |  Generative AI  |  "
    "AWS Bedrock (Claude)  |  REST API & SDK Design  |  API Integration  |  Multi-Tenant "
    "Architecture  |  Full-Stack (Django, Next.js, Vue)  |  Application Security  |  Cloud & "
    "DevOps (AWS, CI/CD)"
)

heading("Professional Experience")
role("Software Engineer — WebAlive", "Jan 2026 – Present")
bullet(
    "Built **Cortex AI**, a multi-tenant AI agent platform: a three-tier Large Language Model "
    "(LLM) routing engine (confidence, action, and knowledge tiers) on AWS Bedrock (Anthropic "
    "Claude), a Retrieval-Augmented Generation (RAG) pipeline with vector search, a Django / "
    "Django REST Framework backend and Vue dashboard, and a plugin protocol — delivered in a "
    "focused six-week sprint and deployed on AWS ECS (Fargate)."
)
bullet(
    "Led a comprehensive **security-hardening** initiative: Argon2 password hashing, split-key "
    "JWT with revocation, SSRF and CORS allowlisting, HttpOnly cookies with CSRF protection, "
    "role-based access control (RBAC), Content-Security-Policy, PII redaction and retention "
    "controls, per-tenant rate limiting, IAM-scoped model access, and RAG prompt-injection "
    "mitigation."
)
bullet(
    "Built **Store Migration AI**, a Model Context Protocol (MCP)-based Shopify-to-WebCommander "
    "migration pipeline with machine-learning-driven API matching, and drove API "
    "standardization for AI-agent and MCP consumption."
)
bullet(
    "Lead developer on a **clinical conversational AI portal** (Next.js): session and survey "
    "lifecycle, a form scenario engine, and a two-tier chatbot lifecycle; reviewed the RAG "
    "ingestion pipeline for a clinical-genetics knowledge assistant."
)

role("Software Engineer — Bitmascot", "Jun 2024 – Dec 2025")
bullet(
    "Architected and shipped **6 production REST API SDKs** (Python, PHP, C# / .NET) for "
    "subscription-commerce, events, and e-commerce platforms — 25+ entity modules each with "
    "full CRUD, OAuth 2.0 authentication, thread-safe token refresh, pagination, and API "
    "versioning."
)
bullet(
    "Built **E.L.I Bot**, an enterprise customer-support chatbot, with an LLM fine-tuning "
    "pipeline, a custom text-chunking algorithm, training UX/UI, and a document-to-training-"
    "data flow."
)
bullet(
    "Delivered core backend APIs (PHP) for a subscription-commerce platform and **tech-led a "
    "C# / .NET SDK**, setting the architecture for HTTP-client management, dual authentication, "
    "and API versioning, and reviewing the full codebase."
)

role("Founder & Lead Developer, Medview  ·  Co-Founder & CTO, Project Durjo", "2022 – 2024")
bullet(
    "Founded two **government-funded (Bangladesh ICT Division)** startups: a medicine-management "
    "platform with OCR prescription reading, and a disaster-management robotics system — led "
    "teams and technical strategy."
)

heading("Projects")
role("Dorpon — Agent-Ready Eyewear Commerce", "")
p = para("**Live demo:** dorpon-eyewear.netlify.app   |   Code: github.com/Meharaj002")
p.paragraph_format.space_after = Pt(2)
bullet(
    "Full-stack premium eyewear store (Next.js, Django / DRF) with browser-based **virtual "
    "try-on** and an AI shopping concierge on AWS Bedrock (Anthropic Claude)."
)
bullet(
    "Exposed the store to AI agents through a stateless **Model Context Protocol (MCP)** server "
    "and a **Universal Commerce Protocol (UCP)** REST layer (discovery manifest, JSON-Schema "
    "actions, RFC 9457 errors), aligned with the Agentic Commerce Protocol — one shared tool "
    "registry powering the web app, MCP, and REST. Deployed on free-tier cloud infrastructure "
    "(Render, Netlify, Neon, Supabase)."
)

heading("Technical Skills")
bullet(
    "**AI & Machine Learning:** Agentic AI, Model Context Protocol (MCP), Large Language Models "
    "(LLMs), Retrieval-Augmented Generation (RAG) — chunking, embeddings, vector databases, LLM "
    "fine-tuning, prompt engineering, Generative AI, Natural Language Processing (NLP), AWS "
    "Bedrock (Anthropic Claude), prompt-injection mitigation; PyTorch, TensorFlow, Keras"
)
bullet("**Languages:** Python, PHP, JavaScript, TypeScript, C# / .NET, Java, C/C++, SQL")
bullet(
    "**Frameworks & Libraries:** Django, Django REST Framework (DRF), async DRF (adrf), "
    "Next.js, Vue.js, React, Vitest"
)
bullet("**Databases & Caching:** PostgreSQL, Redis, DynamoDB, MySQL, SQLite, vector databases")
bullet(
    "**Cloud & DevOps:** AWS (ECS / Fargate, Bedrock, IAM, RDS), CI/CD, Git, GitHub, Bitbucket, "
    "Git LFS"
)
bullet(
    "**Security:** OAuth 2.0, JWT, Argon2, CSRF / CORS / SSRF protection, RBAC, "
    "Content-Security-Policy, rate limiting, PII handling"
)
bullet(
    "**Practices & Tools:** REST API design, SDK design, API integration, microservices, "
    "multi-tenant architecture, Agile, Jira, Confluence"
)

heading("Education & Credentials")
bullet("**B.Sc. in Computer Science** — BRAC University  (2020 – 2024)")
bullet(
    "**Certifications:** Python Programming · Python Data Structures — University of Michigan;  "
    "Social Media Marketing — Meta"
)
bullet(
    "**Award:** Runner-Up, Project Showcase — 5th International Tech Carnival, Bangladesh ICT "
    "Division"
)

doc.save("cv.docx")
print("wrote cv.docx")
